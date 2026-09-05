#!/usr/bin/env python3
"""Generate the facilitator tutorial and example-course documents."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TUTORIAL_DIR = ROOT / "public" / "tutorials"
EXAMPLES_DIR = ROOT / "public" / "examples"
TMP_DIR = ROOT / "tmp" / "tutorial-assets"

NAVY = "0B2E4F"
TEAL = "177E78"
MINT = "E9F6F4"
GOLD = "EAB83E"
PALE_GOLD = "FFF7DD"
INK = "172A3A"
SLATE = "5F7383"
PALE = "F4F8FA"
WHITE = "FFFFFF"
RED = "9B3B35"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=130, bottom=110, end=130) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_fixed_table_widths(table, widths_cm: list[float], indent_twips: int = 130) -> None:
    """Write matching table, grid and cell widths for stable Word/PDF rendering."""
    widths_twips = [round(value / 2.54 * 1440) for value in widths_cm]
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_twips))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, width in enumerate(widths_twips):
            cell = row.cells[index]
            cell.width = Cm(widths_cm[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("UCC Facilitator Resource  •  ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(SLATE)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def setup_document(title: str, subject: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Title", 28, NAVY),
        ("Heading 1", 20, NAVY),
        ("Heading 2", 14, TEAL),
        ("Heading 3", 11, NAVY),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(4)

    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "University of Cape Coast"
    props.keywords = "UCC, microcredential, facilitator, course design, tutorial"
    props.comments = "Illustrative facilitator resource generated for the UCC Microcredential Platform."

    header = section.header
    p = header.paragraphs[0]
    p.text = "UNIVERSITY OF CAPE COAST  /  MICROCREDENTIALS"
    p.style = styles["Caption"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].font.name = "Aptos"
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    p.paragraph_format.space_after = Pt(0)
    add_page_number(section.footer.paragraphs[0])
    return doc


def add_label(doc: Document, text: str, color=TEAL) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.letter_spacing = Pt(0.8) if hasattr(run.font, "letter_spacing") else None


def add_banner(doc: Document, title: str, body: str, fill=MINT, accent=TEAL) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.15)
    p.paragraph_format.right_indent = Cm(0.15)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "22")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r.font.size = Pt(10)
    r = p.add_run(body)
    r.font.size = Pt(9)


def add_bullets(doc: Document, items: list[str], style="List Bullet") -> None:
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_after = Pt(2)
        lead = p.add_run(f"{index}.  ")
        lead.bold = True
        lead.font.color.rgb = RGBColor.from_string(TEAL)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        shade(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        r.font.size = Pt(8.5)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cell = cells[i]
            shade(cell, PALE if row_index % 2 else WHITE)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(value)
            for r in p.runs:
                r.font.size = Pt(8.25)
    if widths:
        set_fixed_table_widths(table, widths)
    else:
        set_fixed_table_widths(table, [15.8 / len(headers)] * len(headers))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_cover(doc: Document, kicker: str, title: str, subtitle: str, version: str) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    add_label(doc, kicker, GOLD)
    p = doc.add_paragraph(style="Title")
    p.add_run(title)
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph(subtitle)
    p.paragraph_format.space_after = Pt(18)
    for r in p.runs:
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor.from_string(SLATE)
    add_banner(
        doc,
        "A worked model—not a prescribed course",
        "Every field is filled to show the expected level of clarity. Facilitators should adapt the evidence, content, dates and assessment to their own discipline and obtain academic approval before delivery.",
        PALE_GOLD,
        GOLD,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(28)
    p = doc.add_paragraph(version)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(TEAL)


def page(doc: Document) -> None:
    doc.add_page_break()


def make_guide(path: Path) -> None:
    doc = setup_document(
        "Facilitator Course Studio Annotated Guide",
        "Six-stage guide and complete illustrative course for the UCC Microcredential Platform",
    )
    add_cover(
        doc,
        "FACILITATOR COURSE STUDIO • ANNOTATED GUIDE",
        "Build a review-ready UCC microcredential",
        "A six-stage tutorial with completion indicators, field-by-field examples and a complete illustrative course.",
        "Release 10 • 5 September 2026",
    )
    doc.add_heading("What this guide gives you", level=2)
    add_bullets(doc, [
        "A clean route through all six Course Studio stages.",
        "Plain-language explanations of required fields and optional enhancements.",
        "A complete model course that can be loaded safely into the unsaved editor.",
        "Checks for outcome alignment, accessibility, authentic evidence and the certificate gate.",
        "A final pre-submission checklist for academic approval.",
    ])
    add_banner(doc, "Start in the application", "Open Facilitator Portal → Course Studio → Load complete example. The example changes only the current unsaved editor state until you choose Save draft.")

    page(doc)
    add_label(doc, "HOW THE STUDIO WORKS")
    doc.add_heading("One course, six visible stages", level=1)
    doc.add_paragraph("The redesigned workspace separates planning, content, practice, assessment and final review. The right-side checks show what is complete and what must be corrected before the draft can be submitted.")
    add_table(doc, ["Stage", "Purpose", "Completion evidence"], [
        ["1 · Essentials", "Define the offer and audience.", "Code, title, description, discipline, level, language, delivery, hours, enrolment and price."],
        ["2 · Outcomes", "Connect purpose, capability and evidence.", "Objectives, measurable outcomes, skill, assessment method, skills tags and course sections."],
        ["3 · Learning", "Create the learner journey.", "Accessible learning blocks mapped to a section and at least one outcome."],
        ["4 · Activities", "Add authentic practice.", "Colab or virtual-lab instructions, rubric, pass mark, attempts, due date and required/optional status."],
        ["5 · Assessment", "Configure valid judgement.", "Assessment modes, scored questions, feedback, marking schemes, pass mark and attempts."],
        ["6 · Review", "Test and submit with confidence.", "Readiness checks, learner preview, saved draft and submission for academic approval."],
    ], [2.7, 5.1, 8.0])
    doc.add_heading("Read the indicators", level=2)
    add_table(doc, ["Indicator", "Meaning", "Action"], [
        ["Green tick", "Every required check for the stage is satisfied.", "Continue, but still review quality and accuracy."],
        ["Clock / incomplete", "A named required field or evidence item is missing.", "Open the stage and follow the exact side-check message."],
        ["Optional", "The enhancement may be omitted.", "If added, configure it fully and ensure it serves an outcome."],
        ["Readiness score", "A design-quality signal, not academic approval.", "Resolve all blocking checks before submission."],
    ], [3.0, 7.0, 5.8])
    add_banner(doc, "Side-check rule", "A completed field is not the same as a good field. Use specific language, observable verbs and evidence that a reviewer can inspect.", PALE_GOLD, GOLD)

    page(doc)
    add_label(doc, "STAGES 1–2")
    doc.add_heading("Define the offer, then align the design", level=1)
    doc.add_heading("Stage 1 · Essentials", level=2)
    add_table(doc, ["Field", "Illustrative entry", "Why it works"], [
        ["Course title", "Applied Data Literacy for Evidence-Based Decision-Making", "Names the capability and practical purpose."],
        ["Code", "UCC-MC-ADL-001", "Unique, short and institutionally traceable."],
        ["Audience", "Public-sector, NGO, academic and business professionals who interpret data.", "Identifies who will use the learning and in what context."],
        ["Prerequisites", "Basic computer and spreadsheet skills; no programming experience required.", "Sets fair expectations without hidden barriers."],
        ["Duration / mode", "30 hours · blended · applied level", "States learner workload and delivery pattern."],
        ["Enrolment / price", "Application · GHS 450", "Makes access and commercial conditions explicit."],
        ["Accessibility", "Structured HTML, heading-led documents, reviewed transcripts and keyboard support.", "Describes concrete provision, not a generic promise."],
    ], [3.0, 7.5, 5.3])
    doc.add_heading("Stage 2 · Outcome-led design", level=2)
    doc.add_paragraph("Write objectives as broad intentions. Write outcomes as observable learner performances and name the evidence that will prove achievement.")
    add_table(doc, ["Outcome statement", "Skill", "Assessment method"], [
        ["Explain the data lifecycle and identify quality, provenance, privacy and ethical risks.", "Conceptual understanding", "Objective knowledge check"],
        ["Clean, summarise and visualise a small dataset in Google Colab using a reproducible workflow.", "Data literacy", "Applied practical evidence"],
        ["Interpret a chart accurately and communicate a finding with uncertainty and limitations.", "Professional communication", "Presentation / demonstration"],
        ["Defend an ethical, evidence-informed recommendation for a realistic decision.", "Ethical decision-making", "Oral assessment"],
    ], [8.8, 3.3, 3.7])

    page(doc)
    add_label(doc, "STAGE 3")
    doc.add_heading("Build an accessible learner journey", level=1)
    doc.add_paragraph("The example uses four sections and seven learning blocks. Every block belongs to a section, has a clear learner-facing title, names its source and maps to one or more outcomes.")
    add_table(doc, ["Section", "Learning block", "Format / design note"], [
        ["1 · Orientation", "Start here: pathway, support and certificate conditions", "Readable HTML · states pathway and credential rules."],
        ["1 · Orientation", "Annotated community-indicators case pack", "PDF download · includes data dictionary and decision prompts."],
        ["2 · Foundations", "Data quality, provenance and ethical use", "Readable HTML · five pre-analysis checks."],
        ["2 · Foundations", "How a neural network learns", "Reviewed external video · link plus companion transcript."],
        ["3 · Applied analysis", "Worked example: community-service dataset", "Readable HTML · question-to-chart workflow."],
        ["3 · Applied analysis", "World Bank Open Data explorer", "External embed/link · source terms checked before publication."],
        ["4 · Capstone", "Capstone evidence and oral-defence brief", "Editable Word plus equivalent readable HTML."],
    ], [4.1, 7.2, 4.5])
    doc.add_heading("Choose the right content method", level=2)
    add_table(doc, ["Use", "When", "Minimum check"], [
        ["Embed readable text", "The content is course-authored or can be lawfully adapted.", "Headings, lists, links and plain-language structure render correctly."],
        ["Upload PDF / Word", "Learners need a fixed or editable reference file.", "Retain a readable HTML equivalent; never display raw binary text."],
        ["Link / embed", "A stable authoritative external source adds value.", "Open in a new tab, name the source and provide an alternative if embedding fails."],
        ["Video", "Motion or demonstration serves an outcome.", "Review captions/transcript, duration, relevance and access conditions."],
    ], [3.6, 7.1, 5.1])
    add_banner(doc, "Document preview safeguard", "The learner preview must render extracted readable text or offer the original download. PDF or Word bytes must never be decoded directly into the lesson window.")

    page(doc)
    add_label(doc, "STAGES 4–5")
    doc.add_heading("Require authentic practice and valid evidence", level=1)
    doc.add_heading("Stage 4 · Programme activities", level=2)
    add_table(doc, ["Activity", "Gate", "Configuration"], [
        ["Colab data-cleaning and visualisation notebook", "Required", "Pass 70% · 3 attempts · instructions, notebook and 100-mark rubric supplied."],
        ["Virtual measurement and data-quality practical", "Required", "Pass 70% · 2 attempts · safety gate, observations, calculation and report rubric."],
        ["Ohm’s law data interpretation", "Optional", "Pass 60% · 3 attempts · enrichment only; not part of the certificate gate."],
    ], [7.0, 2.2, 6.6])
    doc.add_heading("Colab preview: two supported routes", level=3)
    add_numbered(doc, [
        "A valid GitHub .ipynb or Google Colab URL opens directly in a new tab.",
        "An uploaded or bundled notebook opens Colab and downloads the .ipynb file; in Colab choose File → Upload notebook.",
        "If nothing opens, allow pop-ups for the UCC platform and try Open activity again.",
    ])
    doc.add_heading("Stage 5 · Assessment", level=2)
    doc.add_paragraph("The example demonstrates the platform’s complete question repertoire. Use only formats that provide valid evidence for the mapped outcome.")
    add_table(doc, ["Assessment modes (8)", "Question formats demonstrated (13)"], [[
        "Objective quiz; applied assignment; practical evidence; presentation; demonstration; portfolio; oral assessment; reflective evidence",
        "Multiple choice; true/false; fill in; matching; drag and drop; picture matching; video question; short answer; essay; scenario response; oral-defence prompt; evidence-upload prompt; practical assignment",
    ]], [7.8, 8.0])
    add_bullets(doc, [
        "Every scored item includes points, a marking scheme, correct/incorrect feedback and learner advice.",
        "Every question maps to at least one outcome.",
        "The course pass mark is 70%; three attempts are allowed.",
        "Required practical evidence is checked separately from the objective assessment score.",
    ])

    page(doc)
    add_label(doc, "STAGE 6")
    doc.add_heading("Review, preview and submit", level=1)
    doc.add_paragraph("Use Review as a quality-assurance workspace. It should show the learner sequence, outcome mappings, assessment evidence and the credential gate in one place.")
    add_table(doc, ["Review check", "Evidence of completion"], [
        ["Identity and catalogue", "Course code and title are final; audience, prerequisites, price and workload are accurate."],
        ["Alignment", "Every outcome has a skill and assessment method; learning and assessment evidence map back to outcomes."],
        ["Learner journey", "Sections are ordered; titles are meaningful; duration is realistic; links and downloads open."],
        ["Accessibility", "HTML alternatives, headings, transcripts, labels and keyboard operation are checked."],
        ["Activities", "Instructions, rubric, pass mark, attempts, due date and gate status are complete."],
        ["Assessment", "Questions, solutions, feedback and marking schemes have been tested in learner preview."],
        ["Credential", "Certificate is enabled only when the course has an approved, defensible completion rule."],
        ["Governance", "Save draft, preview as student, then submit for academic review; publication remains an administrator action."],
    ], [5.0, 10.8])
    add_banner(doc, "Certificate gate", "Verified identity + at least 70% on the course assessment + completion of every required programme activity = automatic University of Cape Coast digital certificate with a live verification QR.", PALE_GOLD, GOLD)
    doc.add_heading("Before selecting Submit for review", level=2)
    add_bullets(doc, [
        "Replace all illustrative titles, codes, dates, people, datasets and links.",
        "Confirm copyright, licence, privacy and consent for every external or uploaded resource.",
        "Preview every lesson and activity as a student on desktop and mobile.",
        "Ask another facilitator to complete the instructions without verbal explanation.",
        "Resolve every blocking side check and record any justified accessibility alternative.",
    ])

    page(doc)
    add_label(doc, "COMPLETE ILLUSTRATIVE COURSE")
    doc.add_heading("Applied Data Literacy for Evidence-Based Decision-Making", level=1)
    doc.add_paragraph("This model is intentionally complete so facilitators can study the relationship between purpose, curriculum, authentic practice, assessment and credentialing.")
    add_table(doc, ["Design element", "Illustrative specification"], [
        ["Discipline / level", "Professional · Applied"],
        ["Delivery / language", "Blended · English"],
        ["Workload / access", "30 hours · application enrolment · GHS 450"],
        ["Objectives / outcomes / skills", "4 objectives · 4 measurable outcomes · 8 skills tags"],
        ["Curriculum", "4 sections · 7 learning blocks"],
        ["Practice", "3 activities: two required and one optional"],
        ["Assessment", "8 modes · 13 question types · 70% pass mark · 3 attempts"],
        ["Credential", "UCC digital certificate with QR verification after all gate conditions are met"],
    ], [5.4, 10.4])
    doc.add_heading("Four course objectives", level=2)
    add_numbered(doc, [
        "Build practical confidence in asking answerable questions and judging data quality, provenance and limitations.",
        "Develop the ability to clean, summarise and visualise a small dataset with transparent, reproducible methods.",
        "Enable ethical communication of evidence-based recommendations to academic, workplace and community audiences.",
        "Strengthen authentic assessment through process evidence, reflection and a short oral defence.",
    ])
    doc.add_heading("Eight skill tags", level=2)
    doc.add_paragraph("Critical thinking • Digital literacy • Evidence-based decision-making • Applied problem-solving • Communication • Data literacy • Research and inquiry • Responsible data storytelling")
    add_banner(doc, "Use the assets", "The ZIP includes a PDF case pack, editable Word capstone brief, Colab notebook and CSV question bank. They are examples for study and must be adapted before a live course is published.")

    page(doc)
    add_label(doc, "FACILITATOR SELF-CHECK")
    doc.add_heading("Course submission checklist", level=1)
    checklist_rows = [
        ["☐", "The title, code and description clearly distinguish this course."],
        ["☐", "Audience, prerequisites, workload, delivery, enrolment and price are accurate."],
        ["☐", "Objectives describe purpose; outcomes use observable verbs and name evidence."],
        ["☐", "Every learning block belongs to a section and maps to an outcome."],
        ["☐", "PDF/Word items have a readable alternative and external links have been tested."],
        ["☐", "Videos have reviewed captions or a transcript and all images are meaningfully labelled."],
        ["☐", "Each programme activity has instructions, rubric, pass mark, attempts and gate status."],
        ["☐", "Questions provide enough evidence and include marking guidance and feedback."],
        ["☐", "The certificate gate matches the approved achievement standard."],
        ["☐", "Student preview works on desktop/mobile, including downloads, Colab and virtual labs."],
        ["☐", "All illustrative data and placeholder dates have been replaced."],
        ["☐", "A colleague has reviewed the full learner journey before submission."],
    ]
    add_table(doc, ["Done", "Quality assurance statement"], checklist_rows, [1.5, 14.3])
    doc.add_heading("Files supplied with this tutorial", level=2)
    add_bullets(doc, [
        "UCC-Facilitator-Course-Studio-Annotated-Guide.pdf and .docx",
        "Applied-Data-Literacy-Case-Pack.pdf",
        "Applied-Data-Literacy-Capstone-Brief.docx",
        "Applied-Data-Literacy-Colab-Notebook.ipynb",
        "Applied-Data-Literacy-Question-Bank.csv",
    ])
    add_banner(doc, "Institutional note", "The platform can generate a credential only after the configured evidence is complete. Academic approval and the authority to issue a University of Cape Coast credential remain governed by University policy.", PALE_GOLD, GOLD)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def make_capstone(path: Path) -> None:
    doc = setup_document(
        "Applied Data Literacy Capstone Brief",
        "Illustrative capstone evidence and oral-defence brief",
    )
    add_cover(
        doc,
        "ILLUSTRATIVE COURSE • CAPSTONE ASSESSMENT",
        "Evidence-based decision brief",
        "Applied Data Literacy for Evidence-Based Decision-Making",
        "Editable learner brief • Illustrative content",
    )
    doc.add_heading("The challenge", level=2)
    doc.add_paragraph("Use the supplied community-indicators case pack or a facilitator-approved dataset to answer one practical decision question for a named audience. Your submission must show how you moved from source data to a responsible recommendation.")
    add_banner(doc, "Required submission", "One evidence portfolio containing the cleaned dataset, reproducible notebook, accessible chart, 500-word decision brief, cleaning log, references, reflection and oral verification.")
    doc.add_heading("Choose one decision question", level=2)
    add_bullets(doc, [
        "Which learner-support channel should receive additional attention next month?",
        "Which locality or service group needs further investigation before resources are allocated?",
        "What change can be recommended from the evidence without overstating what the data proves?",
    ])
    doc.add_heading("Evidence checklist", level=2)
    add_table(doc, ["Evidence", "Minimum requirement"], [
        ["Question and audience", "A specific, answerable question and the person/group who will use the result."],
        ["Source and permission", "Data source, date, licence/permission, unit of analysis and data dictionary."],
        ["Cleaning log", "Original values preserved; missing, duplicate, range and category checks documented."],
        ["Reproducible analysis", "A working .ipynb notebook with labelled steps and outputs."],
        ["Accessible visual", "Title, units, source, readable labels and a concise text description."],
        ["Decision brief", "Finding, uncertainty, at least two limitations and a proportionate recommendation."],
        ["Reflection", "150–250 words on ethics, bias, privacy and what you would verify next."],
        ["Oral defence", "Three-minute explanation or approved accessible equivalent responding to questions."],
    ], [4.3, 11.5])

    page(doc)
    add_label(doc, "MARKING AND SUBMISSION")
    doc.add_heading("100-mark rubric", level=1)
    add_table(doc, ["Criterion", "Marks", "What strong evidence shows"], [
        ["Question, source and provenance", "15", "Clear decision need; traceable source; appropriate permission and definitions."],
        ["Inspection and cleaning", "20", "Systematic checks, transparent decisions and preserved original data."],
        ["Analysis and reproducibility", "20", "Correct method; notebook runs in order; outputs are labelled and interpretable."],
        ["Visualisation and interpretation", "15", "Suitable chart; accessible labelling; accurate finding and uncertainty."],
        ["Recommendation and limitations", "15", "Evidence supports the action; limitations constrain the claim appropriately."],
        ["Ethics, reflection and citation", "10", "Privacy, bias, consent and source use are considered and referenced."],
        ["Oral defence / verification", "5", "Learner can explain key choices and respond consistently with the submission."],
    ], [5.0, 1.8, 9.0])
    add_banner(doc, "Pass and credential rule", "A minimum assessment result of 70% is required. The UCC digital certificate is issued only when identity is verified and every activity marked required has also been completed.", PALE_GOLD, GOLD)
    doc.add_heading("Submission instructions", level=2)
    add_numbered(doc, [
        "Name files with your learner ID and a short project title; do not include unnecessary personal data.",
        "Upload the completed .ipynb and supporting evidence, or submit an authorised sharing link where enabled.",
        "Check that the notebook opens, all cells run and the chart has an accessible description.",
        "Submit before the course deadline. Up to three assessment attempts are available unless academic approval specifies otherwise.",
        "Keep a personal copy of every submitted file and the confirmation receipt.",
    ])
    doc.add_heading("Academic integrity and responsible use", level=2)
    doc.add_paragraph("You may use approved tools to support learning, but the submitted reasoning, verification and oral defence must be your own. Cite all data, code and external assistance. Do not upload confidential, personal or restricted data to public services.")
    add_banner(doc, "Illustrative resource", "Your facilitator will confirm the live dataset, due date, support route and approved submission method. This file demonstrates the level of detail expected in Course Studio.")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def make_case_pack(path: Path) -> None:
    doc = setup_document(
        "Applied Data Literacy Case Pack",
        "Fictional community-service dataset and data dictionary",
    )
    add_cover(
        doc,
        "ILLUSTRATIVE COURSE • CASE PACK",
        "Community learning-support indicators",
        "A small fictional dataset for data-quality, interpretation and decision-making practice.",
        "Teaching example • Not operational UCC data",
    )
    add_banner(doc, "Safe practice data", "All records and values in this pack are fictional. They must not be interpreted as University of Cape Coast performance data or used for operational decisions.", PALE_GOLD, GOLD)
    doc.add_heading("Scenario", level=2)
    doc.add_paragraph("A microcredential team wants to decide which learner-support channel should receive additional attention next month. Six weekly summaries have been supplied. You must inspect data quality, compare demand and service performance, and recommend an action without claiming more than the dataset supports.")
    doc.add_heading("Fictional weekly data", level=2)
    add_table(doc, ["Week", "Channel", "Requests", "Resolved", "Median min", "Satisfaction /5"], [
        ["1", "Email", "84", "70", "310", "3.7"],
        ["2", "Live chat", "102", "93", "18", "4.3"],
        ["3", "Telephone", "61", "52", "42", "4.0"],
        ["4", "Email", "97", "76", "355", "3.5"],
        ["5", "Live Chat", "116", "108", "17", "4.4"],
        ["6", "Telephone", "58", "—", "46", "3.8"],
    ], [1.8, 3.2, 2.3, 2.3, 3.0, 3.2])
    doc.add_paragraph("Notice the spelling variation in Live chat / Live Chat and the missing resolved value in Week 6. Do not silently repair either issue; record the rule and consequence in your cleaning log.")

    page(doc)
    add_label(doc, "DATA DICTIONARY AND TASKS")
    doc.add_heading("Data dictionary", level=1)
    add_table(doc, ["Variable", "Definition", "Type / valid values"], [
        ["week", "Sequential reporting week in this teaching extract.", "Integer 1–6"],
        ["channel", "Primary support route used by learners.", "Email, Live chat, Telephone"],
        ["requests", "New support requests logged during the week.", "Non-negative integer"],
        ["resolved", "Requests marked resolved by the reporting cut-off.", "Non-negative integer; may be missing"],
        ["median_minutes", "Median elapsed minutes to first substantive response.", "Positive number; minutes"],
        ["satisfaction", "Mean response to a five-point post-service item.", "Number 1.0–5.0"],
    ], [3.4, 7.4, 5.0])
    doc.add_heading("Required analysis", level=2)
    add_numbered(doc, [
        "State the decision question, unit of analysis and limits of the six-row extract.",
        "Check missing values, category consistency, valid ranges and whether resolved can exceed requests.",
        "Create a transparent cleaned version while retaining the original data.",
        "Calculate resolution rate only where the inputs are available; state how missing data is handled.",
        "Create one suitable chart comparing a meaningful indicator by channel.",
        "Write one finding, at least two limitations and a proportionate recommendation.",
    ])
    doc.add_heading("Questions for critical interpretation", level=2)
    add_bullets(doc, [
        "Does high request volume necessarily indicate poor service? What other explanations are plausible?",
        "Can satisfaction scores be compared without knowing the response rate and question wording?",
        "Could requests opened in one week be resolved in a later week? How would that affect a resolution rate?",
        "What disaggregation or contextual evidence is needed before resources are reallocated?",
        "What personal or sensitive data should be excluded from a public notebook?",
    ])
    add_banner(doc, "Facilitator annotation", "A strong learner response cleans only what can be justified, shows all assumptions, distinguishes observation from interpretation and recommends further verification where the evidence is weak.")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> None:
    TUTORIAL_DIR.mkdir(parents=True, exist_ok=True)
    EXAMPLES_DIR.mkdir(parents=True, exist_ok=True)
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    make_guide(TUTORIAL_DIR / "UCC-Facilitator-Course-Studio-Annotated-Guide.docx")
    make_capstone(EXAMPLES_DIR / "Applied-Data-Literacy-Capstone-Brief.docx")
    make_case_pack(TMP_DIR / "Applied-Data-Literacy-Case-Pack.docx")
    print("Generated facilitator tutorial, capstone brief and case pack source.")


if __name__ == "__main__":
    main()
